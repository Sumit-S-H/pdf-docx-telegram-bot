import logging
import os
import threading
from http.server import BaseHTTPRequestHandler, HTTPServer

from telegram import Update
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
    ConversationHandler
)
from fpdf import FPDF
import io
from docx import Document
from docx.shared import Pt, Inches
from collections import defaultdict

# ────────────────────────────────────────────────
# Fetch token from Render environment variables
BOT_TOKEN = os.environ.get("BOT_TOKEN", "")
# ────────────────────────────────────────────────

logging.basicConfig(level=logging.INFO)

# States for conversation after /generate
CHOOSING_FORMAT = 0

# Data per user
user_pages = defaultdict(list)          # user_id → [ [page1 lines], [page2 lines], ... ]
user_current_page = defaultdict(list)   # user_id → current page lines being collected

def normalize_marker(text: str) -> int | None:
    t = text.strip().lower()
    if t in ("p1", "page 1", "1"):
        return 1
    if t.startswith(("p", "page ")):
        try:
            cleaned = t.replace("p", "").replace("page ", "").strip()
            num = int(cleaned)
            if num >= 1:
                return num
        except:
            pass
    return None

# ─── PDF GENERATOR ────────────────────────────────────────
def generate_pdf(pages: list[list[str]]) -> io.BytesIO:
    pdf = FPDF()
    pdf.set_font("Arial", size=12)
    pdf.set_margins(20, 20, 20)
    LINE_HEIGHT = 6

    for page_lines in pages:
        pdf.add_page()
        for line in page_lines:
            # This line removes emojis and special characters that the PDF font doesn't support
            clean_line = line.encode('latin-1', 'ignore').decode('latin-1')
            
            if not clean_line.strip():
                pdf.ln(LINE_HEIGHT)
                continue
            pdf.multi_cell(170, LINE_HEIGHT, clean_line, align="L")
            pdf.ln(3)  # spacing between paragraphs

    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

# ─── DOCX GENERATOR ───────────────────────────────────────
def generate_docx(pages: list[list[str]]) -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    for i, page_lines in enumerate(pages, 1):
        if i > 1:
            doc.add_page_break()

        for line in page_lines:
            if line.strip():
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(8)
            else:
                doc.add_paragraph()  # empty line

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    text = update.message.text.strip()

    if not text:
        return

    marker = normalize_marker(text)

    if marker is not None:
        # Close current page if it has content
        if user_current_page[user_id]:
            user_pages[user_id].append(user_current_page[user_id][:])
            user_current_page[user_id].clear()

        # Fill skipped pages with empty ones if user jumps numbers
        while len(user_pages[user_id]) < marker - 1:
            user_pages[user_id].append([])

        await update.message.reply_text(
            f"✅ Page {marker} finished.\n"
            f"Now collecting content for page {marker + 1}.\n"
            "When ready → /generate"
        )
    else:
        # Normal paragraph
        user_current_page[user_id].append(text)


async def cmd_generate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id

    # Finish last page if user didn't send final pX
    if user_current_page[user_id]:
        user_pages[user_id].append(user_current_page[user_id][:])
        user_current_page[user_id].clear()

    if not user_pages[user_id]:
        await update.message.reply_text("No content yet. Send some text first!")
        return ConversationHandler.END

    await update.message.reply_text(
        "Which format do you want?\n\n"
        "Reply with: **PDF** or  **DOCX**"
    )

    return CHOOSING_FORMAT


async def handle_format_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    choice = update.message.text.strip().lower()

    if choice in ("pdf", "docx"):
        await update.message.reply_text(f"Generating {choice.upper()}... 📄")

        try:
            if choice == "pdf":
                file_buffer = generate_pdf(user_pages[user_id])
                filename = "document.pdf"
            else:
                file_buffer = generate_docx(user_pages[user_id])
                filename = "document.docx"

            await update.message.reply_document(
                document=file_buffer,
                filename=filename,
                caption=f"Your document ({len(user_pages[user_id])} page{'s' if len(user_pages[user_id]) > 1 else ''}) – {choice.upper()}"
            )

        except Exception as e:
            await update.message.reply_text(f"Error while creating file: {str(e)}")

        return ConversationHandler.END

    else:
        await update.message.reply_text("Please reply with **PDF** or **DOCX** only.")
        return CHOOSING_FORMAT


async def cmd_clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    user_pages[user_id].clear()
    user_current_page[user_id].clear()
    await update.message.reply_text("🗑️ Everything cleared. Ready for a new document.")


async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    help_text = (
        "📄 **Welcome to Text → PDF/DOCX Bot!**\n\n"
        "This bot turns your messages into a beautiful multi-page PDF or Word document.\n\n"
        "How to use it (very simple):\n\n"
        "1. Just send normal text messages — each message becomes a **paragraph** on the current page.\n"
        "   You can send as many as you want.\n\n"
        "2. When you want to finish the current page and start a new one, send:\n"
        "   • `p1`   or   `page 1`   or   `1`     → ends page 1, starts page 2\n"
        "   • `p2`   or   `page 2`   or   `2`     → ends page 2, starts page 3\n"
        "   • and so on...\n\n"
        "3. When you've added all your content, send:\n"
        "   `/generate`   (or `/gerneater` if you prefer 😉)\n\n"
        "4. The bot will ask: **PDF** or **DOCX**?\n"
        "   Just reply with `pdf` or `docx` (not case-sensitive)\n\n"
        "5. You'll receive your multi-page document ready to download!\n\n"
        "Extra commands:\n"
        "• `/clear`   → delete everything and start a new document\n"
        "• `/start`   → show this help message again\n\n"
        "Ready when you are — start sending text! ✍️"
    )

    await update.message.reply_text(
        help_text,
        parse_mode="Markdown",
        disable_web_page_preview=True
    )

# ─── DUMMY WEB SERVER FOR RENDER ──────────────────────────
class DummyHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.send_header('Content-type', 'text/plain')
        self.end_headers()
        self.wfile.write(b"Bot is running successfully!")
    
    def log_message(self, format, *args):
        # Suppress noisy HTTP logs in your terminal
        pass

def run_dummy_server():
    # Render automatically sets the PORT environment variable
    port = int(os.environ.get("PORT", 8080))
    server = HTTPServer(('0.0.0.0', port), DummyHandler)
    print(f"Dummy server listening on port {port} for Render health checks...")
    server.serve_forever()
# ──────────────────────────────────────────────────────────

def main():
    # Start the dummy server in a background thread so it doesn't block the bot
    threading.Thread(target=run_dummy_server, daemon=True).start()

    if not BOT_TOKEN:
        print("ERROR: BOT_TOKEN is not set in environment variables!")
        return

    app = Application.builder().token(BOT_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("generate", cmd_generate),
                      CommandHandler("gerneater", cmd_generate)],
        states={
            CHOOSING_FORMAT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_format_choice)
            ]
        },
        fallbacks=[],
    )

    app.add_handler(conv_handler)
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("clear", cmd_clear))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    print("Telegram Bot polling started...")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
